import streamlit as st
import asyncio
import sqlite3
from datetime import datetime
import json
import os
from io import BytesIO
import time
import re
import requests
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from groq import Groq
from config.settings import SETTINGS
import urllib.parse
import google.generativeai as genai
import base64
import traceback


# =========================
# Helper: Pollinations AI Image Generation
# =========================

# =========================
# Helper: Gemini Image Generation (ONLY)
# =========================

def get_gemini_model():
    """
    Return a Gemini image model using either env var or Streamlit Settings.
    """
    api_key = (
        os.environ.get("GEMINI_API_KEY")
        or st.session_state.get("gemini_api_key")
    )
    if not api_key:
        st.warning("⚠️ Gemini API key not set. Please configure it in Settings.")
        return None

    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("models/gemini-2.5-flash-image")
        return model
    except Exception as e:
        st.error(f"Failed to initialize Gemini model: {e}")
        return None


def extract_gemini_image_bytes(response):
    """
    Extract image bytes from Gemini response (same logic as your working script).
    """
    try:
        if not hasattr(response, "candidates"):
            return None

        for c in response.candidates:
            content = getattr(c, "content", None)
            if not content:
                continue

            parts = getattr(content, "parts", None) or []
            for p in parts:
                # Inline image data (SDK object)
                if hasattr(p, "inline_data") and getattr(p.inline_data, "data", None):
                    return p.inline_data.data

                # Dict-like fallback
                if isinstance(p, dict):
                    if "inline_data" in p and "data" in p["inline_data"]:
                        return p["inline_data"]["data"]

        return None
    except Exception:
        traceback.print_exc()
        return None


def generate_gemini_image(prompt: str) -> str | None:
    """
    Generate an image with Gemini and return a data URL (data:image/png;base64,...)
    so it works in Streamlit, HTML, PDF, DOCX.
    """
    model = get_gemini_model()
    if not model:
        return None

    try:
        # EXACTLY like your working script: no response_mime_type, just prompt
        response = model.generate_content(prompt)
        image_bytes = extract_gemini_image_bytes(response)

        if not image_bytes:
            st.error("❌ No image data found in Gemini response.")
            return None

        # Convert to base64 data URL
        encoded = base64.b64encode(image_bytes).decode("utf-8")
        data_url = f"data:image/png;base64,{encoded}"
        return data_url

    except Exception as e:
        st.error(f"Gemini Image Error: {e}")
        traceback.print_exc()
        return None


def fetch_pexels_images(query: str, num_images: int = 3, api_key: str | None = None):
    """
    Fetch image URLs from Pexels API based on query.
    Returns a list of image URLs (landscape images).
    """
    if not api_key:
        return []

    try:
        url = "https://api.pexels.com/v1/search"
        headers = {"Authorization": api_key}
        params = {"query": query, "per_page": num_images, "orientation": "landscape"}
        resp = requests.get(url, headers=headers, params=params, timeout=10)

        if resp.status_code != 200:
            return []

        data = resp.json()
        photos = data.get("photos", [])
        image_urls = [p["src"]["large"] for p in photos if "src" in p and "large" in p["src"]]
        return image_urls
    except Exception:
        return []


# =========================
# Helper: Groq LLM Integration
# =========================

def get_groq_client():
    """
    Return a Groq client using either env var or session state key.
    """
    api_key = (
        os.environ.get("GROQ_API_KEY")
        or st.session_state.get("groq_api_key")
        or getattr(SETTINGS, "GROQ_API_KEY", None)
    )
    if not api_key:
        return None, "Groq API key not found. Please set GROQ_API_KEY env var or add it in Settings."
    try:
        client = Groq(api_key=api_key)
        return client, None
    except Exception as e:
        return None, f"Failed to initialize Groq client: {e}"


def generate_blog_with_llm(
    topic: str,
    category: str,
    niche: str,
    keywords: str,
    target_audience: str,
    content_intent: str,
    expertise_level: str,
    tone: str,
    length: str,
    writing_style: str,
    additional_context: str = "",
):
    """
    Use Groq (Llama-3.1) to generate a full blog in markdown-style text.
    """
    client, err = get_groq_client()
    if err:
        raise RuntimeError(err)

    length_hint = {
        "Short (800-1000 words)": "around 800-1000 words",
        "Medium (2000-2500 words)": "around 2000-2500 words",
        "Long (3500-4000 words)": "around 3500-4000 words",
        "Comprehensive (5000-6000 words)": "around 5000-6000 words",
    }.get(length, "around 3000 words")

    prompt = f"""
You are an expert content writer and SEO strategist.

Write a DETAILED, COMPREHENSIVE blog article in well-structured markdown with:

- H1 title
- H2 / H3 headings
- Bullet lists, numbered lists
- No separator lines like "====" or "----" or "___" or "****"

Topic: {topic}
Category: {category}
Niche: {niche or "Not specified"}
Primary Keywords: {keywords or "Not specified"}
Audience: {target_audience}
Intent: {content_intent}
Expertise Level: {expertise_level}
Tone: {tone}
Writing Style: {writing_style}
Target Length: {length_hint}

Additional Context:
{additional_context or "No extra context."}
"""

    response = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[
            {
                "role": "system",
                "content": "You are a world-class blog writer and SEO expert. "
                           "Use clean markdown headings, bold/italic, lists. "
                           "Avoid using separator-only lines like ====, ----, ****, ___."
            },
            {"role": "user", "content": prompt},
        ],
        temperature=0.7,
        max_tokens=8000,
    )

    content = response.choices[0].message.content
    return content


def normalize_markdown_with_llm(raw_content: str, topic: str) -> str:
    """
    SECOND PASS (Option D):
    Take raw blog text and convert it into clean markdown with correct
    # / ## / ### headings and bullet lists, without changing wording much.
    """
    client, err = get_groq_client()
    if err or not raw_content.strip():
        return raw_content

    norm_prompt = f"""
You are a markdown formatter.

Convert the following blog article about "{topic}" into CLEAN MARKDOWN.

Requirements:
- Preserve the original wording as much as possible.
- Detect logical section titles and mark them as markdown headings:
  - Main sections -> '## Title'
  - Subsections -> '### Title'
- Keep the existing order of content.
- Use '-' for bullet lists and '1.' for numbered lists.
- Do NOT add extra commentary or explanations.
- Do NOT wrap the result in ``` fences.
Return ONLY the converted markdown.
"""

    try:
        response = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {"role": "system", "content": "You convert plain text articles into clean markdown."},
                {"role": "user", "content": norm_prompt + "\n\nArticle:\n" + raw_content},
            ],
            temperature=0.2,
            max_tokens=8000,
        )
        formatted = response.choices[0].message.content.strip()
        # Safety: if it's too short or empty, fall back
        if len(formatted.split()) < len(raw_content.split()) * 0.4:
            return raw_content
        return formatted
    except Exception:
        return raw_content


def generate_image_prompts(topic: str, num_prompts: int = 3):
    """
    Generate relevant image prompts based on the blog topic using LLM.
    """
    client, err = get_groq_client()
    if err:
        return []

    prompt = f"""
Generate {num_prompts} highly specific and detailed image generation prompts for a blog about "{topic}".

CRITICAL:
- Directly related to {topic}
- 20-30 words each
- Include visual style keywords

Return ONLY a JSON array of {num_prompts} strings.
Example: ["prompt 1", "prompt 2", "prompt 3"]
"""

    try:
        response = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {"role": "system", "content": "You are an expert at creating image generation prompts."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.8,
            max_tokens=500,
        )

        content = response.choices[0].message.content.strip()
        prompts = json.loads(content)
        return prompts if isinstance(prompts, list) else []
    except Exception as e:
        st.warning(f"Could not generate image prompts: {e}")
        return []


def generate_image_captions_from_prompts(topic: str, prompts):
    """
    Convert detailed prompts into short figure captions (max 8 words).
    Captions must be concise, clean, and human-friendly.
    """

    if not prompts:
        return []

    client, err = get_groq_client()
    if err:
        # Fallback: simple shortening
        captions = []
        for p in prompts:
            short = re.sub(r"[^\w\s]", "", p).strip().split()
            captions.append(" ".join(short[:8]))   # limit to 8 words
        return captions

    # LLM Prompt
    system_prompt = "You create short, concise figure captions."
    user_prompt = f"""
Generate a caption for each prompt.

Rules for each caption:
- Maximum 8 words
- Human-friendly
- No technical jargon
- No style words (e.g., cinematic, 8k, ultra HD)
- No verbs like 'illustrating', 'showing', 'depicting'
- Should describe the figure in simple terms

Return ONLY a JSON array of captions.

Prompts:
{json.dumps(prompts, indent=2)}
"""

    try:
        resp = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3,
            max_tokens=300,
        )

        captions = json.loads(resp.choices[0].message.content.strip())

        # Final cleaning + enforce 8-word rule
        clean_caps = []
        for c in captions:
            c = re.sub(r"[^\w\s]", "", c).strip()
            clean_caps.append(" ".join(c.split()[:8]))

        return clean_caps

    except Exception:
        # fallback
        fallback = []
        for p in prompts:
            c = re.sub(r"[^\w\s]", "", p).strip().split()
            fallback.append(" ".join(c[:8]))
        return fallback



# =========================
# Text Cleaning Helpers
# =========================

def clean_text(text: str) -> str:
    """Remove emojis & unsupported chars for PDF safety."""
    return re.sub(r"[^\x00-\x7F]+", " ", text)


def clean_separator_lines(text: str) -> str:
    """
    Remove separator lines like ====, ----, ___, ****, #######.
    """
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if stripped and all(c in "=-_*#" for c in stripped):
            continue
        cleaned.append(line)
    return "\n".join(cleaned)


def clean_markdown_formatting(text: str) -> str:
    """
    Remove markdown formatting like **bold**, *italic*, and leading # from headings.
    """
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)([^*]+?)\*(?!\*)", r"\1", text)
    text = re.sub(r"^\s*#{1,6}\s*", "", text)
    return text

def sanitize_caption(caption: str, max_words: int = 8):
    # remove symbols
    caption = re.sub(r"[^\w\s]", " ", caption)
    caption = re.sub(r"\s+", " ", caption).strip()

    # limit to 8 words always
    words = caption.split()
    caption = " ".join(words[:max_words])

    # final fallback: ensure not empty & not a single long word
    if len(caption) == 0:
        caption = "Image"

    # break long unbreakable sequences
    if len(max(words, key=len)) > 20:
        caption = " ".join([w[:20] for w in words[:max_words]])

    return caption

# =========================
# PDF Generation
# =========================

def add_image_to_pdf(pdf: FPDF, img_url_or_path: str, max_width: int = 160, caption: str = ""):
    try:
        # Load image
        if img_url_or_path.startswith("data:image"):
            # Handle data URL from Gemini
            header, b64data = img_url_or_path.split(",", 1)
            img = Image.open(BytesIO(base64.b64decode(b64data)))
        elif img_url_or_path.startswith("http"):
            headers = {"User-Agent": "Mozilla/5.0"}
            resp = requests.get(img_url_or_path, timeout=20, stream=True, headers=headers)
            resp.raise_for_status()
            img = Image.open(BytesIO(resp.content))
        else:
            img = Image.open(img_url_or_path)


        if img.mode in ("RGBA", "LA", "P"):
            bg = Image.new("RGB", img.size, (255, 255, 255))
            img = img.convert("RGBA")
            bg.paste(img, mask=img.split()[3])
            img = bg

        w, h = img.size
        aspect = h / w

        usable_width = pdf.w - pdf.l_margin - pdf.r_margin
        new_w = min(max_width, usable_width)
        new_h = new_w * aspect

        # Convert px → mm (3.78 px/mm)
        resized = img.resize((int(new_w * 3.78), int(new_h * 3.78)), Image.Resampling.LANCZOS)

        temp = f"temp_{int(time.time()*1000)}.jpg"
        resized.save(temp, "JPEG", quality=85)

        x = pdf.l_margin + (usable_width - new_w) / 2
        pdf.image(temp, x=x, w=new_w)

        os.remove(temp)

        pdf.ln(4)

        if caption:
            safe_caption = sanitize_caption(caption)
            page_width = pdf.w - pdf.l_margin - pdf.r_margin
            pdf.set_font("helvetica", "I", 9)
            pdf.set_text_color(110, 110, 110)
            pdf.multi_cell(page_width, 5, safe_caption, 0, "C")
            pdf.set_text_color(40, 40, 40)
            pdf.ln(3)

        return True

    except Exception as e:
        pdf.set_font("helvetica", "I", 9)
        pdf.set_text_color(150, 150, 150)
        pdf.multi_cell(0, 5, "[Image could not be loaded]", 0, "C")
        pdf.set_text_color(40, 40, 40)
        pdf.ln(4)
        return False

def sanitize_caption(caption: str, max_words: int = 8):
    caption = re.sub(r"[^\w\s]", " ", caption)
    caption = re.sub(r"\s+", " ", caption).strip()

    words = caption.split()
    if not words:
        return "Image"

    # limit caption to 8 words
    words = words[:max_words]

    # ensure no extremely long token (FPDF breaks here)
    words = [w[:20] for w in words]

    return " ".join(words)


def generate_pdf(title, content, meta_info, images=None, image_descriptions=None):
    images = images or []
    image_descriptions = image_descriptions or []

    content = clean_separator_lines(content)
    content = clean_text(content)
    title = title[:150]

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_left_margin(15)
    pdf.set_right_margin(15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # --- Title ---
    pdf.set_font("helvetica", "B", 20)
    pdf.multi_cell(0, 10, clean_markdown_formatting(title), 0, "C")
    pdf.ln(4)

    # Divider
    pdf.set_draw_color(52, 152, 219)
    pdf.set_line_width(0.8)
    y = pdf.get_y()
    pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
    pdf.ln(8)

    # Meta Info
    pdf.set_font("helvetica", "", 10)
    pdf.set_text_color(90, 90, 90)
    pdf.cell(0, 5, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", 0, 1)

    for key, val in meta_info.items():
        label = key.replace("_", " ").title()
        pdf.cell(0, 5, f"{label}: {val}", 0, 1)

    pdf.set_text_color(40, 40, 40)
    pdf.ln(8)

    # Images
    if images:
        pdf.set_font("helvetica", "B", 14)
        pdf.cell(0, 8, "Visual Content", 0, 1)
        pdf.ln(2)

        for i, img in enumerate(images):
            desc = image_descriptions[i] if i < len(image_descriptions) else "Image"
            caption = f"Figure {i+1}: {sanitize_caption(desc)}"
            add_image_to_pdf(pdf, img, max_width=170, caption=caption)

        pdf.ln(4)

    # --- Content Rendering ---
    heading_pattern = re.compile(r'^\s*(#{1,6})\s*(.*)')
    page_width = pdf.w - pdf.l_margin - pdf.r_margin

    for raw in content.split("\n"):
        line = raw.strip()

        if not line:
            pdf.ln(3)
            continue

        # --- Headings ---
        m = heading_pattern.match(line)
        if m:
            level = len(m.group(1))
            heading = clean_markdown_formatting(m.group(2))

            if level == 1:
                pdf.set_font("helvetica", "B", 16)
            elif level == 2:
                pdf.set_font("helvetica", "B", 13)
            else:
                pdf.set_font("helvetica", "B", 12)

            pdf.multi_cell(0, 7, heading, 0, "L")
            pdf.ln(2)
            continue

        # --- Numbered list: 1. Text ---
        num = re.match(r"^(\d+)\.\s+(.*)", line)
        if num:
            number = num.group(1)
            text = clean_markdown_formatting(num.group(2))

            pdf.set_font("helvetica", "", 11)

            indent = pdf.l_margin + 5
            pdf.set_x(indent)
            pdf.cell(8, 6, f"{number}.", 0, 0)

            pdf.multi_cell(page_width - 13, 6, text, 0, "L")
            continue

        # --- Bullet list ---
        if line.startswith("- ") or line.startswith("* "):
            pdf.set_font("helvetica", "", 11)
            bullet = clean_markdown_formatting(line[2:])
            pdf.set_x(pdf.l_margin + 5)
            pdf.multi_cell(page_width - 5, 6, f"- {bullet}", 0, "L")
            continue

        # --- Blockquote ---
        if line.startswith("> "):
            pdf.set_font("helvetica", "I", 11)
            pdf.set_text_color(120, 120, 120)
            quote = clean_markdown_formatting(line[2:])
            pdf.multi_cell(page_width, 6, quote, 0, "L")
            pdf.set_text_color(40, 40, 40)
            continue

        # --- Normal text ---
        pdf.set_font("helvetica", "", 11)
        safe = re.sub(r"[^\x00-\x7F]+", " ", line)
        pdf.multi_cell(page_width, 6, safe, 0, "L")

    return bytes(pdf.output())



# =========================
# DOCX Generation
# =========================

def generate_docx(title, content, meta_info, images=None, image_descriptions=None):
    """
    Generate DOCX report with enhanced formatting using cleaned markdown.
    """
    images = images or []
    image_descriptions = image_descriptions or []

    content = clean_separator_lines(content)

    doc = Document()

    # Title
    title_para = doc.add_heading(clean_markdown_formatting(title), 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Meta info
    meta_para = doc.add_paragraph()
    run = meta_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.italic = True

    run = meta_para.add_run(f"Word Count: {meta_info.get('word_count', 'N/A')}\n")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    run = meta_para.add_run(f"Reading Time: {meta_info.get('reading_time', 'N/A')}\n")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # Images
    if images:
        doc.add_heading("Visual Content", level=2)
        doc.add_paragraph()

        for idx, img_url in enumerate(images, 1):
            try:
                if img_url.startswith("data:image"):
                    # Data URL from Gemini
                    header, b64data = img_url.split(",", 1)
                    img = Image.open(BytesIO(base64.b64decode(b64data)))

                    tmp_file = f"temp_docx_img_{int(time.time() * 1000)}_{idx}.jpg"
                    img.save(tmp_file, "JPEG", quality=85)

                    doc.add_picture(tmp_file, width=Inches(6))

                    caption_para = doc.add_paragraph()
                    desc = image_descriptions[idx-1] if idx-1 < len(image_descriptions) else "Content Image"
                    caption_run = caption_para.add_run(f"Figure {idx}: {desc}")
                    caption_run.font.size = Pt(9)
                    caption_run.italic = True
                    caption_run.font.color.rgb = RGBColor(100, 100, 100)
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    try:
                        os.remove(tmp_file)
                    except Exception:
                        pass

                elif img_url.startswith("http"):
                    # existing HTTP branch (unchanged)
                    headers = {
                        "User-Agent": "Mozilla/5.0",
                        "Accept": "image/*",
                    }
                    resp = requests.get(img_url, timeout=25, stream=True, headers=headers)
                    resp.raise_for_status()

                    img = Image.open(BytesIO(resp.content))
                    if img.mode in ("RGBA", "LA", "P"):
                        background = Image.new("RGB", img.size, (255, 255, 255))
                        if img.mode == "P":
                            img = img.convert("RGBA")
                        background.paste(img, mask=img.split()[-1] if img.mode == "RGBA" else None)
                        img = background

                    tmp_file = f"temp_docx_img_{int(time.time() * 1000)}_{idx}.jpg"
                    img.save(tmp_file, "JPEG", quality=85)
                    doc.add_picture(tmp_file, width=Inches(6))

                    caption_para = doc.add_paragraph()
                    desc = image_descriptions[idx-1] if idx-1 < len(image_descriptions) else "Content Image"
                    caption_run = caption_para.add_run(f"Figure {idx}: {desc}")
                    caption_run.font.size = Pt(9)
                    caption_run.italic = True
                    caption_run.font.color.rgb = RGBColor(100, 100, 100)
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    try:
                        os.remove(tmp_file)
                    except Exception:
                        pass

                else:
                    # local path
                    doc.add_picture(img_url, width=Inches(6))
                    caption_para = doc.add_paragraph()
                    desc = image_descriptions[idx-1] if idx-1 < len(image_descriptions) else "Content Image"
                    caption_run = caption_para.add_run(f"Figure {idx}: {desc}")
                    caption_run.font.size = Pt(9)
                    caption_run.italic = True
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


                doc.add_paragraph()
            except Exception:
                p = doc.add_paragraph()
                desc = image_descriptions[idx-1] if idx-1 < len(image_descriptions) else "Content Image"
                run = p.add_run(f"[Figure {idx}: {desc} could not be loaded]")
                run.italic = True
                run.font.color.rgb = RGBColor(150, 150, 150)

    heading_pattern = re.compile(r'^\s*(#{1,6})\s*(.*)')

    # Content
    for line in content.split("\n"):
        line_stripped = line.strip()
        if not line_stripped:
            doc.add_paragraph()
            continue

        m = heading_pattern.match(line_stripped)
        if m:
            level = len(m.group(1))
            heading_text = clean_markdown_formatting(line_stripped)
            if level == 1:
                doc.add_heading(heading_text, level=1)
            elif level == 2:
                doc.add_heading(heading_text, level=2)
            elif level == 3:
                doc.add_heading(heading_text, level=3)
            else:
                doc.add_heading(heading_text, level=4)
            continue

        if line_stripped.startswith("- ") or line_stripped.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(clean_markdown_formatting(line_stripped[2:]))
            continue

        # Label-style paragraph (e.g., "Benefits of X:")
        label, sep, rest = line_stripped.partition(":")
        if sep and len(label.split()) <= 5 and all(w[:1].isupper() for w in label.split() if w):
            para = doc.add_paragraph()
            run_label = para.add_run(clean_markdown_formatting(label) + ":")
            run_label.bold = True
            if rest.strip():
                para.add_run(" " + clean_markdown_formatting(rest.strip()))
        else:
            para = doc.add_paragraph()
            para.add_run(clean_markdown_formatting(line_stripped))

    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return docx_file.getvalue()


# =========================
# HTML Generation
# =========================

def generate_html(title, content, meta_info, images=None, image_descriptions=None):
    """
    Generate styled HTML report, cleaning markdown markers so **, ****, #### etc
    do NOT appear in the final HTML.
    """
    images = images or []
    image_descriptions = image_descriptions or []

    title_clean = clean_markdown_formatting(title)
    content = clean_separator_lines(content)

    html_template = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title_clean}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.8;
            max-width: 900px;
            margin: 0 auto;
            padding: 30px;
            color: #333;
            background-color: #f8f9fa;
        }}
        .container {{
            background: white;
            padding: 40px;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 4px solid #3498db;
            padding-bottom: 15px;
            margin-bottom: 30px;
            font-size: 2.5em;
        }}
        h2 {{
            color: #34495e;
            margin-top: 40px;
            margin-bottom: 20px;
            font-size: 1.8em;
            border-left: 4px solid #3498db;
            padding-left: 15px;
        }}
        h3 {{
            color: #555;
            margin-top: 30px;
            margin-bottom: 15px;
            font-size: 1.4em;
        }}
        h4, h5, h6 {{
            color: #555;
            margin-top: 25px;
            margin-bottom: 10px;
        }}
        .meta-info {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 20px;
            border-radius: 8px;
            margin: 25px 0;
            font-size: 0.95em;
        }}
        .meta-info strong {{
            color: #fff;
            font-weight: 600;
        }}
        .image-gallery {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 20px;
            margin: 30px 0;
        }}
        .image-gallery img {{
            width: 100%;
            height: auto;
            border-radius: 10px;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
            transition: transform 0.3s ease;
        }}
        .image-gallery img:hover {{
            transform: scale(1.02);
        }}
        p {{
            margin-bottom: 15px;
            line-height: 1.8;
        }}
        blockquote {{
            border-left: 5px solid #3498db;
            margin: 25px 0;
            padding: 15px 25px;
            color: #555;
            font-style: italic;
            background: #ecf0f1;
            border-radius: 5px;
        }}
        ul, ol {{
            margin: 20px 0;
            padding-left: 30px;
        }}
        li {{
            margin-bottom: 12px;
            line-height: 1.6;
        }}
        .content {{
            margin-top: 30px;
        }}
        .caption {{
            text-align: center;
            font-size: 0.85em;
            color: #666;
            margin-top: 5px;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>{title_clean}</h1>

        <div class="meta-info">
            <p><strong>📅 Generated:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
            <p><strong>📝 Word Count:</strong> {meta_info.get('word_count', 'N/A')}</p>
            <p><strong>⏱️ Reading Time:</strong> {meta_info.get('reading_time', 'N/A')}</p>
            <p><strong>🎯 SEO Score:</strong> {meta_info.get('seo_score', 'N/A')}</p>
        </div>
"""

    if images:
        html_template += "        <h2>Visual Content</h2>\n"
        html_template += '        <div class="image-gallery">\n'
        for i, img_url in enumerate(images):
            desc = image_descriptions[i] if i < len(image_descriptions) else "Content Image"
            html_template += "            <div>\n"
            html_template += f'                <img src="{img_url}" alt="Content image {i+1}" loading="lazy" />\n'
            html_template += f'                <div class="caption">Figure {i+1}: {desc}</div>\n'
            html_template += "            </div>\n"
        html_template += "        </div>\n"

    html_template += '        <div class="content">\n'

    heading_pattern = re.compile(r'^\s*(#{1,6})\s*(.*)')
    in_list = False

    for line in content.split("\n"):
        raw = line.rstrip()
        stripped = raw.strip()

        if not stripped:
            if in_list:
                html_template += "        </ul>\n"
                in_list = False
            html_template += "        <br>\n"
            continue

        m = heading_pattern.match(stripped)
        if m:
            if in_list:
                html_template += "        </ul>\n"
                in_list = False

            level = len(m.group(1))
            heading_text = clean_markdown_formatting(stripped)

            if level == 1:
                html_template += f"        <h1>{heading_text}</h1>\n"
            elif level == 2:
                html_template += f"        <h2>{heading_text}</h2>\n"
            elif level == 3:
                html_template += f"        <h3>{heading_text}</h3>\n"
            elif level == 4:
                html_template += f"        <h4>{heading_text}</h4>\n"
            elif level == 5:
                html_template += f"        <h5>{heading_text}</h5>\n"
            else:
                html_template += f"        <h6>{heading_text}</h6>\n"
            continue

        if stripped.startswith("> "):
            if in_list:
                html_template += "        </ul>\n"
                in_list = False
            html_template += f"        <blockquote>{clean_markdown_formatting(stripped[2:])}</blockquote>\n"
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            if not in_list:
                html_template += "        <ul>\n"
                in_list = True
            text = clean_markdown_formatting(stripped[2:])
            html_template += f"            <li>{text}</li>\n"
            continue

        if in_list:
            html_template += "        </ul>\n"
            in_list = False

        label, sep, rest = stripped.partition(":")
        if sep and len(label.split()) <= 5 and all(w[:1].isupper() for w in label.split() if w):
            label_html = clean_markdown_formatting(label)
            rest_html = clean_markdown_formatting(rest.strip()) if rest.strip() else ""
            html_template += f"        <p><strong>{label_html}:</strong> {rest_html}</p>\n"
        else:
            text = clean_markdown_formatting(stripped)
            html_template += f"        <p>{text}</p>\n"

    if in_list:
        html_template += "        </ul>\n"

    html_template += """
        </div>
    </div>
</body>
</html>
"""
    return html_template.encode("utf-8")


# =========================
# STREAMLIT APP
# =========================

st.set_page_config(
    page_title="End-to-End Blog Creation Suite",
    page_icon="📝",
    layout="wide",
)

st.title("📝 End-to-End Blog Creation Suite")
st.markdown("Complete automated blog creation with AI-powered content and image generation")

st.sidebar.title("Navigation")
page = st.sidebar.selectbox(
    "Select Page",
    ["Dashboard", "Blog Creation", "Agent Monitor", "Analytics", "Settings"],
)

# -------------- Dashboard --------------
if page == "Dashboard":
    st.header("📊 Dashboard")

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Active Agents", "25", "↑2")
    with col2:
        st.metric("Blogs Created", "147", "↑12")
    with col3:
        st.metric("Success Rate", "94.2%", "↑1.2%")
    with col4:
        st.metric("Avg. Quality Score", "8.7/10", "↑0.3")

    st.subheader("Recent Activity")
    activity_data = [
        {"Time": "10:30 AM", "Agent": "Title Generation Agent", "Status": "✅ Completed", "Blog": "AI Trends 2025"},
        {"Time": "10:28 AM", "Agent": "SEO Optimization Agent", "Status": "🔄 In Progress", "Blog": "Tech Review"},
        {"Time": "10:25 AM", "Agent": "Research Agent", "Status": "✅ Completed", "Blog": "Market Analysis"},
    ]
    st.table(activity_data)

# -------------- Blog Creation --------------
elif page == "Blog Creation":
    st.header("✍️ Create New Blog")
    st.markdown("Configure your blog with comprehensive options")

    with st.form("blog_creation_form"):
        st.subheader("🎯 Core Topic Information")
        col1, col2 = st.columns(2)
        with col1:
            topic = st.text_input(
                "Blog Topic *",
                placeholder="e.g., Artificial Intelligence in Healthcare",
            )
            category = st.selectbox(
                "Category",
                [
                    "Technology",
                    "Business",
                    "Health & Wellness",
                    "Finance",
                    "Education",
                    "Marketing",
                    "Lifestyle",
                    "Science",
                    "Politics",
                    "Other",
                ],
            )
        with col2:
            niche = st.text_input(
                "Specific Niche",
                placeholder="e.g., AI diagnostics, Telemedicine",
            )
            keywords = st.text_input(
                "Primary Keywords (comma-separated)",
                placeholder="e.g., AI healthcare, medical diagnosis, patient care",
            )

        st.subheader("🎯 Audience & Content Intent")
        col1, col2, col3 = st.columns(3)
        with col1:
            target_audience = st.selectbox(
                "Target Audience",
                [
                    "General Public",
                    "Industry Professionals",
                    "Business Leaders",
                    "Technical Experts",
                    "Students",
                    "Researchers",
                    "Entrepreneurs",
                ],
            )
        with col2:
            content_intent = st.selectbox(
                "Content Intent",
                [
                    "Informational",
                    "Educational",
                    "Commercial",
                    "Transactional",
                    "Problem-Solving",
                    "Thought Leadership",
                    "News & Updates",
                ],
            )
        with col3:
            expertise_level = st.selectbox(
                "Expertise Level",
                ["Beginner", "Intermediate", "Advanced", "Expert"],
            )

        st.subheader("🎨 Writing Style & Format")
        col1, col2, col3 = st.columns(3)
        with col1:
            tone = st.selectbox(
                "Tone",
                [
                    "Professional",
                    "Conversational",
                    "Authoritative",
                    "Friendly",
                    "Inspirational",
                    "Academic",
                    "Persuasive",
                    "Humorous",
                ],
            )
        with col2:
            length = st.selectbox(
                "Desired Length",
                [
                    "Short (800-1000 words)",
                    "Medium (2000-2500 words)",
                    "Long (3500-4000 words)",
                    "Comprehensive (5000-6000 words)",
                ],
                index=2,
            )
        with col3:
            writing_style = st.selectbox(
                "Writing Style",
                [
                    "How-to Guide",
                    "Listicle",
                    "Case Study",
                    "Opinion Piece",
                    "News Article",
                    "Tutorial",
                    "Review",
                    "Comparison",
                ],
            )

        st.subheader("🖼️ Image Generation Options")
        col1, col2 = st.columns(2)
        with col1:
            use_ai_images = st.checkbox("Generate AI Images (Gemini)", value=True)
            num_ai_images = st.slider("Number of AI Images", 1, 5, 3)
        with col2:
            use_pexels = st.checkbox("Also fetch from Pexels", value=False)
            if use_pexels:
                num_pexels_images = st.slider("Number of Pexels Images", 1, 5, 2)

        st.subheader("💡 Additional Context & Instructions")
        additional_context = st.text_area(
            "Provide any additional context, requirements, or specific points to cover",
            placeholder="e.g., Focus on recent developments, include case studies, avoid technical jargon...",
            height=120,
        )

        submitted = st.form_submit_button("🚀 Create Blog with AI", use_container_width=True)

    if submitted and topic:
        try:
            with st.spinner("🤖 Generating blog content with Llama-3.1..."):
                blog_raw = generate_blog_with_llm(
                    topic=topic,
                    category=category,
                    niche=niche,
                    keywords=keywords,
                    target_audience=target_audience,
                    content_intent=content_intent,
                    expertise_level=expertise_level,
                    tone=tone,
                    length=length,
                    writing_style=writing_style,
                    additional_context=additional_context,
                )

            with st.spinner("🧹 Formatting headings & layout..."):
                blog_content = normalize_markdown_with_llm(blog_raw, topic)

            st.session_state["blog_created"] = True
            st.session_state["blog_topic"] = topic
            st.session_state["blog_title"] = f"{topic}: A Comprehensive Guide"
            st.session_state["blog_content"] = blog_content

            words = len(blog_content.split())
            reading_time_min = max(1, int(words / 230))
            meta_info = {
                "word_count": f"{words}",
                "reading_time": f"{reading_time_min} min",
                "seo_score": "N/A",
                "readability": "N/A",
                "keyword_density": "N/A",
                "plagiarism": "Not checked",
            }
            st.session_state["meta_info"] = meta_info

            # Generate images
            all_images = []
            image_descriptions = []

            if use_ai_images:
                with st.spinner("🎨 Generating AI images with Gemini..."):

                    # 1️⃣ Prefer your structured infographic/flowchart prompts
                    image_prompts = [
                        f"{topic} infographic, clean layout, labeled sections, icons, flat vector design, pastel colors, corporate look, boxes with labels, arrows showing relationships, modern UI infographic, high resolution",
                        f"{topic} overview infographic, blocks with titles, arrows, minimal flow structure, flat modern vector style, pastel theme",
                        f"{topic} visual summary infographic, simplified labeled blocks, icons, modern flat UI, clean corporate infographic style",
                        f"{topic} flowchart diagram, rectangular blocks with labels, arrows between steps, white background, thin lines, flat vector style, professional process diagram",
                        f"{topic} decision flowchart, diamond decision nodes, labeled rectangles, directional arrows, minimal pastel colors, clean schematic diagram",
                        f"{topic} process flowchart, linear step-by-step boxes, arrows connecting each stage, modern vector workflow diagram"
                    ][:num_ai_images]

                    # 2️⃣ Fallback: if something goes wrong computing list
                    if not image_prompts:
                        image_prompts = generate_image_prompts(topic, num_ai_images)

                    # 3️⃣ Generate captions
                    image_titles = generate_image_captions_from_prompts(topic, image_prompts)

                    # 4️⃣ Gemini image generation loop
                    st.info(f"🎨 Generating {len(image_prompts)} AI images with Gemini for: {topic}")

                    for idx, prompt in enumerate(image_prompts):
                        img_url = generate_gemini_image(prompt)

                        if img_url:
                            all_images.append(img_url)
                            desc = image_titles[idx] if idx < len(image_titles) else topic
                            image_descriptions.append(f"{desc} (Source: Gemini)")
                            st.success(f"✓ Generated image {idx + 1}/{len(image_prompts)} using Gemini")
                        else:
                            st.error(f"❌ Gemini failed for image {idx + 1}")

                        time.sleep(0.5)



            if use_pexels:
                with st.spinner("📸 Fetching images from Pexels..."):
                    pexels_key = st.session_state.get("pexels_api_key")
                    if pexels_key:
                        pexels_images = fetch_pexels_images(topic, num_pexels_images, pexels_key)
                        all_images.extend(pexels_images)
                        for _ in pexels_images:
                            image_descriptions.append(f"Stock photo related to {topic}")
                    else:
                        st.warning("Pexels API key not found. Please add it in Settings.")

            st.session_state["blog_images"] = all_images
            st.session_state["image_descriptions"] = image_descriptions

            st.success("🎯 Blog generation completed!")

            progress_bar = st.progress(0)
            status_text = st.empty()
            phases = [
                "Ideation",
                "Research",
                "Content Generation",
                "Image Creation",
                "SEO Optimization",
                "Export Preparation",
            ]
            for i, phase in enumerate(phases):
                status_text.markdown(f"### Phase {i+1}: {phase}")
                time.sleep(0.3)
                progress_bar.progress((i + 1) / len(phases))

            st.balloons()
            st.success("✅ All AI agents completed their tasks!")

        except Exception as e:
            st.error(f"Error during blog generation: {e}")
            import traceback
            st.code(traceback.format_exc())

    # Preview & Export
    if st.session_state.get("blog_created", False):
        st.markdown("---")
        st.header("📄 Generated Blog Preview")

        meta_info = st.session_state.get("meta_info", {})
        blog_content = st.session_state.get("blog_content", "")
        blog_title = st.session_state.get("blog_title", "Blog")
        blog_images = st.session_state.get("blog_images", [])
        image_descriptions = st.session_state.get("image_descriptions", [])

        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Word Count", meta_info.get("word_count", "N/A"))
            st.metric("Reading Time", meta_info.get("reading_time", "N/A"))
        with col2:
            st.metric("SEO Score", meta_info.get("seo_score", "N/A"))
            st.metric("Readability", meta_info.get("readability", "N/A"))
        with col3:
            st.metric("Keyword Density", meta_info.get("keyword_density", "N/A"))
            st.metric("Images Generated", len(blog_images))

        if blog_images:
            st.subheader("🖼️ Generated Images")
            cols = st.columns(min(3, len(blog_images)))
            for idx, img_url in enumerate(blog_images):
                with cols[idx % 3]:
                    try:
                        desc = image_descriptions[idx] if idx < len(image_descriptions) else "Content Image"
                        cap = f"Figure {idx+1}: {desc}"
                        if len(cap) > 80:
                            cap = cap[:77] + "..."
                        st.image(
                            img_url,
                            caption=cap,
                            use_column_width=True,
                        )
                    except Exception:
                        st.warning(f"Could not load image {idx+1}")

        st.subheader("📝 Content Preview")
        with st.expander("View Full Content", expanded=True):
            st.markdown(blog_content)

        st.subheader("💾 Export Options")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        col1, col2, col3 = st.columns(3)

        with col1:
            try:
                with st.spinner("Generating PDF..."):
                    pdf_bytes = generate_pdf(
                        blog_title,
                        blog_content,
                        meta_info,
                        images=blog_images,
                        image_descriptions=image_descriptions,
                    )
                st.download_button(
                    label="📄 Download PDF",
                    data=pdf_bytes,
                    file_name=f"{blog_title.replace(' ', '_')}_{timestamp}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"PDF Error: {str(e)}")

        with col2:
            try:
                with st.spinner("Generating DOCX..."):
                    docx_bytes = generate_docx(
                        blog_title,
                        blog_content,
                        meta_info,
                        images=blog_images,
                        image_descriptions=image_descriptions,
                    )
                st.download_button(
                    label="📝 Download DOCX",
                    data=docx_bytes,
                    file_name=f"{blog_title.replace(' ', '_')}_{timestamp}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"DOCX Error: {str(e)}")

        with col3:
            try:
                html_bytes = generate_html(
                    blog_title,
                    blog_content,
                    meta_info,
                    images=blog_images,
                    image_descriptions=image_descriptions,
                )
                st.download_button(
                    label="🌐 Download HTML",
                    data=html_bytes,
                    file_name=f"{blog_title.replace(' ', '_')}_{timestamp}.html",
                    mime="text/html",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"HTML Error: {str(e)}")

# -------------- Agent Monitor --------------
elif page == "Agent Monitor":
    st.header("🤖 Agent Monitor")

    col1, col2 = st.columns([2, 1])
    with col1:
        st.subheader("Active Agents")
        agent_data = [
            {"Agent": "Content Generation Agent", "Phase": "Core System", "Status": "Active", "Load": 85},
            {"Agent": "Image Generation Agent", "Phase": "Visual", "Status": "Active", "Load": 72},
            {"Agent": "SEO Optimization Agent", "Phase": "SEO", "Status": "Processing", "Load": 91},
            {"Agent": "Format Export Agent", "Phase": "Export", "Status": "Idle", "Load": 15},
            {"Agent": "Quality Check Agent", "Phase": "Review", "Status": "Active", "Load": 68},
        ]
        for agent in agent_data:
            with st.expander(f"{agent['Agent']} - {agent['Status']}"):
                col_a, col_b = st.columns(2)
                with col_a:
                    st.write(f"**Phase:** {agent['Phase']}")
                    st.write(f"**Status:** {agent['Status']}")
                with col_b:
                    st.write(f"**Load:** {agent['Load']}%")
                    st.progress(agent["Load"] / 100)

    with col2:
        st.subheader("System Health")
        st.metric("CPU Usage", "67%")
        st.metric("Memory", "4.2GB/8GB")
        st.metric("Queue Length", "12")
        if st.button("🔄 Refresh Status"):
            st.rerun()

# -------------- Analytics --------------
elif page == "Analytics":
    st.header("📈 Analytics & Insights")

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Blog Performance")
        import pandas as pd
        import numpy as np

        dates = pd.date_range("2024-01-01", periods=30, freq="D")
        views = np.random.randint(100, 1000, 30)
        chart_data = pd.DataFrame({"Date": dates, "Views": views})
        st.line_chart(chart_data.set_index("Date"))

    with col2:
        st.subheader("Agent Efficiency")
        import pandas as pd

        agent_performance = pd.DataFrame(
            {
                "Agent Type": ["Content Generation", "Image Generation", "SEO Optimization", "Export"],
                "Success Rate": [94.2, 92.5, 89.7, 97.1],
            }
        )
        st.bar_chart(agent_performance.set_index("Agent Type"))

    st.subheader("Recent Blog Analytics")
    import pandas as pd

    blog_stats = pd.DataFrame(
        {
            "Blog Title": ["AI Trends 2025", "Tech Innovation Guide", "Market Analysis Report"],
            "Views": [2341, 1876, 1543],
            "Engagement": [8.7, 7.2, 9.1],
            "Images": [3, 5, 2],
        }
    )
    st.dataframe(blog_stats, use_container_width=True)

# -------------- Settings --------------
elif page == "Settings":
    st.header("⚙️ Settings")

    with st.form("settings_form"):
        st.subheader("API Configuration")

        groq_api = st.text_input(
            "Groq API Key (Required for content generation)",
            type="password",
            value=st.session_state.get("groq_api_key", ""),
            help="Get your API key from console.groq.com",
        )

        pexels_api = st.text_input(
            "Pexels API Key (Optional - for stock photos)",
            type="password",
            value=st.session_state.get("pexels_api_key", ""),
            help="Get your free API key from pexels.com/api",
        )
        gemini_api = st.text_input(
        "Gemini API Key (Required for image generation)",
        type="password",
        value=st.session_state.get("gemini_api_key", ""),
        help="Get your API key from https://aistudio.google.com"
        )

        st.info("ℹ️ Gemini 2.5 Flash Image is used for all AI image generation.")


        st.subheader("Default Blog Settings")
        default_tone = st.selectbox("Default Tone", ["Professional", "Conversational", "Authoritative"])
        default_length = st.selectbox("Default Length", ["Medium", "Short", "Long"])

        st.subheader("Image Generation Settings")
        default_num_images = st.slider("Default Number of AI Images", 1, 5, 3)
        image_width = st.slider("Image Width (pixels)", 512, 1920, 1024)
        image_height = st.slider("Image Height (pixels)", 512, 1920, 768)

        st.subheader("Export Settings")
        include_images_pdf = st.checkbox("Include images in PDF", value=True)
        include_images_docx = st.checkbox("Include images in DOCX", value=True)

        save_settings = st.form_submit_button("💾 Save Settings")

        if save_settings:
            st.session_state["pexels_api_key"] = pexels_api
            st.session_state["groq_api_key"] = groq_api
            st.session_state["gemini_api_key"] = gemini_api
            st.session_state["default_num_images"] = default_num_images
            st.session_state["image_width"] = image_width
            st.session_state["image_height"] = image_height
            st.success("✅ Settings saved successfully!")
            st.balloons()

# Footer
st.markdown("---")
st.markdown(
    """
<div style='text-align: center; color: #666;'>
    <p>🤖 Powered by AI | Built with Streamlit, Groq (Llama-3.1) & Gemini AI</p>
    <p>📝 Blog Creation | 🎨 Image Generation | 📄 Multi-format Export</p>
</div>
""",
    unsafe_allow_html=True,
)
